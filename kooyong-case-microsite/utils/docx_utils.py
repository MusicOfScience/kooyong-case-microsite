import os
from datetime import datetime
from docx import Document

GEN_DIR = os.path.join(os.getcwd(), "generated_docs")
os.makedirs(GEN_DIR, exist_ok=True)

def _safe(val): return val if val is not None else ""

def generate_meeting_docx(data: dict) -> str:
    """Generates a simple meeting brief DOCX.
    Later we can map fields to your content-control templates.
    """
    doc = Document()
    doc.add_heading("Meeting Brief", level=1)

    pairs = [
        ("Meeting Title", data.get("meeting_title")),
        ("Location", data.get("location")),
        ("Date and time", data.get("meeting_dt")),
        ("Attendees", data.get("attendees")),
        ("Contact person and phone/email", f"{_safe(data.get('contact_person'))} | {_safe(data.get('contact_phone'))} | {_safe(data.get('contact_email'))}"),
        ("Purpose of meeting", data.get("purpose")),
        ("Previous meetings/engagements", data.get("previous_engagements")),
        ("Who initiated the meeting?", data.get("initiated_by")),
        ("Agenda", data.get("agenda")),
        ("Background notes", data.get("background")),
        ("Notes from meeting", data.get("notes")),
    ]
    for k,v in pairs:
        doc.add_paragraph(f"{k}: {_safe(v)}")

    iso = data.get("timestamp_iso","").replace(":","").replace("-","")
    family = data.get("family_name","anon")
    fname = f"MeetingBrief_{family}_{iso}.docx"
    fpath = os.path.join(GEN_DIR, fname)
    doc.save(fpath)
    return fpath

def generate_event_docx(data: dict) -> str:
    """Generates an event brief DOCX."""
    doc = Document()
    doc.add_heading("Event Brief", level=1)

    pairs = [
        ("Name of event", data.get("event_name")),
        ("Location", data.get("location")),
        ("Purpose of event and MR’s attendance", data.get("purpose")),
        ("Date", data.get("date")),
        ("Start time", data.get("start")),
        ("End time", data.get("end")),
        ("Staff attending", data.get("staff")),
        ("Volunteers attending", data.get("volunteers")),
        ("Key contact", f"{_safe(data.get('contact_name'))} | {_safe(data.get('contact_phone'))} | {_safe(data.get('contact_email'))}"),
        ("Who will meet MR?", data.get("who_meets")),
        ("Notes on other invitees", data.get("other_invitees")),
        ("Previous engagements w/ group", data.get("previous")),
        ("Who initiated invitation?", data.get("initiated_by")),
        ("Carparking details", data.get("parking")),
        ("Speech/official duties", data.get("duties")),
        ("Dress code", data.get("dress")),
        ("Background to Org/Event", data.get("background")),
    ]
    for k,v in pairs:
        doc.add_paragraph(f"{k}: {_safe(v)}")

    iso = data.get("timestamp_iso","").replace(":","").replace("-","")
    family = data.get("family_name","anon")
    fname = f"EventBrief_{family}_{iso}.docx"
    fpath = os.path.join(GEN_DIR, fname)
    doc.save(fpath)
    return fpath
